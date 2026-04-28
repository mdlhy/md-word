"""Generate .docx template files for md-formula-converter."""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")


def set_east_asian_font(style, font_name):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)


def add_sample_lists(doc):
    for i in range(1, 4):
        doc.add_paragraph(f"编号项 {i}", style="List Number")
    for i in range(1, 4):
        doc.add_paragraph(f"要点项 {i}", style="List Bullet")


def create_academic():
    doc = Document()

    # A4, margins: top/bottom=2.54cm, left/right=3.17cm
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    set_east_asian_font(normal, '宋体')
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(24)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    set_east_asian_font(h1, '黑体')
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Pt(0)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    set_east_asian_font(h2, '黑体')
    h2.paragraph_format.first_line_indent = Pt(0)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    set_east_asian_font(h3, '黑体')
    h3.paragraph_format.first_line_indent = Pt(0)

    try:
        quote = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        quote = doc.styles['Quote']
    quote.font.italic = True
    quote.paragraph_format.left_indent = Cm(2)

    try:
        code = doc.styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
    except ValueError:
        code = doc.styles['Code']
    code.font.name = 'Consolas'
    code.font.size = Pt(10)

    add_sample_lists(doc)

    doc.save(os.path.join(TEMPLATES_DIR, 'academic.docx'))
    print("Created templates/academic.docx")


def create_homework():
    doc = Document()

    # A4, margins: 2.54cm all
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(10.5)
    set_east_asian_font(normal, '宋体')
    pf = normal.paragraph_format
    pf.line_spacing = 1.25
    pf.first_line_indent = Pt(21)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(12)
    h1.font.bold = True
    set_east_asian_font(h1, '黑体')
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Pt(0)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(10.5)
    h2.font.bold = True
    set_east_asian_font(h2, '黑体')
    h2.paragraph_format.first_line_indent = Pt(0)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(10.5)
    h3.font.bold = True
    h3.font.italic = True
    set_east_asian_font(h3, '黑体')
    h3.paragraph_format.first_line_indent = Pt(0)

    try:
        quote = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        quote = doc.styles['Quote']
    quote.font.italic = True
    quote.paragraph_format.left_indent = Cm(2)

    try:
        code = doc.styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
    except ValueError:
        code = doc.styles['Code']
    code.font.name = 'Consolas'
    code.font.size = Pt(9)

    add_sample_lists(doc)

    doc.save(os.path.join(TEMPLATES_DIR, 'homework.docx'))
    print("Created templates/homework.docx")


def create_report():
    doc = Document()

    # A4, margins: 2.5cm all
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 微软雅黑 + Calibri, 10.5pt, 1.15x line spacing, no first-line indent
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    set_east_asian_font(normal, '微软雅黑')
    pf = normal.paragraph_format
    pf.line_spacing = 1.15
    pf.first_line_indent = Pt(0)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.bold = True
    set_east_asian_font(h1, '微软雅黑')
    h1.paragraph_format.first_line_indent = Pt(0)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    set_east_asian_font(h2, '微软雅黑')
    h2.paragraph_format.first_line_indent = Pt(0)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    set_east_asian_font(h3, '微软雅黑')
    h3.paragraph_format.first_line_indent = Pt(0)

    try:
        quote = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        quote = doc.styles['Quote']
    quote.font.italic = True
    quote.paragraph_format.left_indent = Cm(1.5)

    try:
        code = doc.styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
    except ValueError:
        code = doc.styles['Code']
    code.font.name = 'Consolas'
    code.font.size = Pt(10)

    add_sample_lists(doc)

    doc.save(os.path.join(TEMPLATES_DIR, 'report.docx'))
    print("Created templates/report.docx")


if __name__ == '__main__':
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    create_academic()
    create_homework()
    create_report()
    print("All templates generated successfully.")
