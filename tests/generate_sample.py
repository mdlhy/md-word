"""Generate a test .docx file with LaTeX formula text for converter testing."""
from docx import Document
from docx.shared import Inches
import os


def generate_sample_docx(output_path: str = "test_sample.docx"):
    doc = Document()

    # === Section 1: Simple inline formulas ===
    doc.add_heading("LaTeX 公式测试文档", level=1)

    p = doc.add_paragraph()
    p.add_run("简单行内公式：")
    p.add_run("$x^2 + y^2 = z^2$")
    p.add_run(" 是勾股定理。")

    # === Section 2: Complex inline formula ===
    p = doc.add_paragraph()
    p.add_run("偏导数：")
    p.add_run("$\\frac{\\partial f}{\\partial x}$")
    p.add_run(" 表示函数 f 对 x 的偏导。")

    # === Section 3: Display math ===
    doc.add_paragraph("质能方程：")
    doc.add_paragraph("$$E = mc^2$$")

    # === Section 4: Matrix ===
    doc.add_paragraph("矩阵公式：")
    doc.add_paragraph("$$\\begin{pmatrix} a & b \\\\ c & d \\end{pmatrix}$$")

    # === Section 5: Aligned ===
    doc.add_paragraph("对齐公式组：")
    doc.add_paragraph("$$\\begin{aligned} x &= 1 \\\\ y &= 2 \\end{aligned}$$")

    # === Section 6: Cases ===
    doc.add_paragraph("分段函数：")
    doc.add_paragraph("$$f(x) = \\begin{cases} 1 & x > 0 \\\\ 0 & x \\leq 0 \\end{cases}$$")

    # === Section 7: Nested fractions ===
    doc.add_paragraph("嵌套分数：")
    doc.add_paragraph("$$\\frac{1}{1 + \\frac{1}{x}}$$")

    # === Section 8: Currency (should NOT be converted) ===
    p = doc.add_paragraph()
    p.add_run("价格是 $100。")

    # === Section 9: \(...\) inline ===
    p = doc.add_paragraph()
    p.add_run("Paren 行内公式：")
    p.add_run("\\(a + b = c\\)")
    p.add_run(" 是加法。")

    # === Section 10: \[...\] display ===
    doc.add_paragraph("Bracket 行间公式：")
    doc.add_paragraph("\\[c^2 = a^2 + b^2\\]")

    # === Section 11: Multi-run formula (simulates Word run splitting) ===
    p = doc.add_paragraph()
    p.add_run("多 run 公式：$")
    p.add_run("x^2")
    p.add_run("$ 的值。")  # Formula split across 3 runs

    # === Section 12: Table with formula ===
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "公式"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "$E=mc^2$"
    table.cell(1, 1).text = "质能方程"

    # === Section 13: Add a placeholder image ===
    try:
        from PIL import Image
        import io
        img = Image.new('RGB', (100, 100), color='white')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        doc.add_picture(img_bytes, width=Inches(1.0))
    except ImportError:
        doc.add_paragraph("[图片占位 - PIL不可用]")

    # === Section 14: Normal text paragraph (no formula) ===
    doc.add_paragraph("这是一段普通文本，不包含任何公式。")

    # Save
    doc.save(output_path)
    print(f"Generated: {output_path}")
    return output_path


if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    output = os.path.join(project_dir, "test_sample.docx")
    generate_sample_docx(output)
