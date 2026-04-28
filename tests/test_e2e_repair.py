"""E2E test: repair broken .docx files."""
import sys
import os
import tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from converter.docx_repair import repair_docx
from converter.templates import TEMPLATES

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")


def test_repair_headings():
    """Test fixing # prefixed headings."""
    path = os.path.join(FIXTURES_DIR, "broken_headings.docx")
    if not os.path.exists(path):
        return [("broken_headings", "SKIP", "File not found")]
    
    results = []
    for template_name in TEMPLATES:
        try:
            doc, report = repair_docx(path, template_name)
            
            headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
            
            hash_in_heading = any(p.text.startswith("#") for p in headings)
            
            results.append((
                f"headings/{template_name}",
                "PASS" if not hash_in_heading and len(headings) > 0 else "FAIL",
                f"headings={len(headings)}, hash_in_heading={hash_in_heading}, items={len(report.items)}"
            ))
        except Exception as e:
            results.append((f"headings/{template_name}", "FAIL", str(e)[:80]))
    
    return results


def test_repair_lists():
    """Test fixing list-like plain text."""
    path = os.path.join(FIXTURES_DIR, "broken_lists.docx")
    if not os.path.exists(path):
        return [("broken_lists", "SKIP", "File not found")]
    
    results = []
    for template_name in TEMPLATES:
        try:
            doc, report = repair_docx(path, template_name)
            
            list_items = [p for p in doc.paragraphs if "List" in p.style.name]
            
            results.append((
                f"lists/{template_name}",
                "PASS" if len(list_items) > 0 else "FAIL",
                f"list_items={len(list_items)}, items={len(report.items)}"
            ))
        except Exception as e:
            results.append((f"lists/{template_name}", "FAIL", str(e)[:80]))
    
    return results


def test_repair_mixed():
    """Test fixing mixed issues."""
    path = os.path.join(FIXTURES_DIR, "broken_mixed.docx")
    if not os.path.exists(path):
        return [("broken_mixed", "SKIP", "File not found")]
    
    results = []
    for template_name in ["academic"]:
        try:
            doc, report = repair_docx(path, template_name)
            
            headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
            list_items = [p for p in doc.paragraphs if "List" in p.style.name]
            
            results.append((
                f"mixed/{template_name}",
                "PASS" if len(report.items) > 0 else "FAIL",
                f"headings={len(headings)}, list_items={len(list_items)}, items={len(report.items)}"
            ))
        except Exception as e:
            results.append((f"mixed/{template_name}", "FAIL", str(e)[:80]))
    
    return results


def main():
    from tests.generate_broken_docx import generate_broken_headings, generate_broken_lists, generate_broken_mixed
    generate_broken_headings()
    generate_broken_lists()
    generate_broken_mixed()
    
    all_results = []
    all_results.extend(test_repair_headings())
    all_results.extend(test_repair_lists())
    all_results.extend(test_repair_mixed())
    
    print("\n" + "=" * 80)
    print("E2E Test Results: .docx Repair")
    print("=" * 80)
    passed = 0
    failed = 0
    for name, status, detail in all_results:
        symbol = "✅" if status == "PASS" else "❌" if status == "FAIL" else "⏭️"
        print(f"{symbol} {name}: {detail}")
        if status == "PASS":
            passed += 1
        elif status == "FAIL":
            failed += 1
    
    print(f"\nTotal: {passed + failed} | Passed: {passed} | Failed: {failed}")
    return failed == 0


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
