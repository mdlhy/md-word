import base64
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest
from docx import Document

from converter.markdown_pipeline import convert_markdown_file_to_docx
from converter.pandoc_driver import is_pandoc_available


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgG"
    "M7m6c9QAAAABJRU5ErkJggg=="
)


def _document_metrics(path):
    doc = Document(str(path))
    ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    return {
        "headings": sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading")),
        "tables": len(doc.tables),
        "math": len(doc.element.body.findall(".//m:oMath", ns)),
        "images": len(doc.inline_shapes),
        "paragraph_text": "\n".join(p.text for p in doc.paragraphs),
    }


@pytest.mark.skipif(not is_pandoc_available(), reason="pandoc not installed")
def test_legacy_and_pandoc_cover_core_document_features(tmp_path):
    image_path = tmp_path / "sample.png"
    image_path.write_bytes(PNG_1X1)
    md_path = tmp_path / "sample.md"
    md_path.write_text(
        "# 一级标题\n\n"
        "段落包含 [链接](https://example.com) 和公式 $x^2 + y^2 = z^2$。\n\n"
        "![示例图](sample.png)\n\n"
        "| 指标 | 数值 |\n"
        "|---|---|\n"
        "| **准确率** | 95% |\n",
        encoding="utf-8",
    )

    legacy_path = tmp_path / "legacy.docx"
    pandoc_path = tmp_path / "pandoc.docx"

    legacy = convert_markdown_file_to_docx(
        str(md_path),
        str(legacy_path),
        template_name="academic",
        engine="legacy",
    )
    pandoc = convert_markdown_file_to_docx(
        str(md_path),
        str(pandoc_path),
        template_name="academic",
        engine="pandoc",
    )

    legacy_metrics = _document_metrics(legacy_path)
    pandoc_metrics = _document_metrics(pandoc_path)

    for metrics in (legacy_metrics, pandoc_metrics):
        assert metrics["headings"] >= 1
        assert metrics["tables"] == 1
        assert metrics["math"] >= 1
        assert metrics["images"] >= 1
        assert "链接" in metrics["paragraph_text"]

    assert legacy.compat_report.summary.get("high", 0) == 0
    assert pandoc.compat_report.summary.get("high", 0) == 0
