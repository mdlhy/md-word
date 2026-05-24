import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest

from converter.md_converter import convert_md_to_docx
from converter.md_parser import parse_markdown


def test_fenced_code_language_is_normalized():
    token = parse_markdown("```python\nprint('hello')\n```")[0]

    assert token.type == "code"
    assert token.attrs["language"] == "python"
    assert token.attrs["lang"] == "python"


def test_inline_link_text_is_not_dropped():
    doc, _ = convert_md_to_docx("阅读 [OpenAI](https://openai.com) 文档")

    assert any("阅读 OpenAI 文档" in p.text for p in doc.paragraphs)


def test_heading_keeps_rich_inline_text():
    doc, _ = convert_md_to_docx("# **Bold** title")
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

    assert any("Bold title" in text for text in headings)


def test_table_keeps_rich_inline_text():
    doc, _ = convert_md_to_docx(
        "| A | B |\n"
        "|---|---|\n"
        "| **bold** | [link](https://example.com) |\n"
    )

    assert len(doc.tables) == 1
    assert doc.tables[0].cell(1, 0).text == "bold"
    assert doc.tables[0].cell(1, 1).text == "link"


@pytest.mark.skipif(
    not pytest.importorskip("converter.diagram_renderer").HAS_MATPLOTLIB,
    reason="matplotlib not installed",
)
def test_diagram_code_block_is_rendered_as_image():
    doc, _ = convert_md_to_docx(
        "```matrix\n"
        "name: A\n"
        "1 2\n"
        "3 4\n"
        "caption: Matrix A\n"
        "```"
    )

    assert len(doc.inline_shapes) >= 1
    assert any("Matrix A" in p.text for p in doc.paragraphs)


def test_math_renderer_fallback_is_used_when_math2docx_fails(monkeypatch):
    import converter.md_converter as md_converter
    from docx.oxml import OxmlElement

    def fail_add_math(*args, **kwargs):
        raise RuntimeError("math2docx failed")

    monkeypatch.setattr(md_converter.math2docx, "add_math", fail_add_math)
    monkeypatch.setattr(
        "converter.math_renderer.latex_to_omml",
        lambda latex: OxmlElement("m:oMath"),
    )

    doc, _ = convert_md_to_docx("$x^2$")
    ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}

    assert doc.element.body.findall(".//m:oMath", ns)


def test_cli_conversion_resolves_relative_image_paths(tmp_path, monkeypatch):
    import converter.md_converter as md_converter
    from converter.cli import convert_single

    md_file = tmp_path / "doc.md"
    image_file = tmp_path / "image.png"
    out_file = tmp_path / "doc.docx"
    image_file.write_bytes(b"not a real image")
    md_file.write_text("![Local image](image.png)", encoding="utf-8")
    seen_sources = []

    def fake_add_image(doc, token, template_config):
        seen_sources.append(token.attrs.get("src"))
        return doc.add_paragraph("image")

    monkeypatch.setattr(md_converter, "add_image", fake_add_image)

    assert convert_single(str(md_file), str(out_file), "academic", False, False, engine="legacy")
    assert seen_sources == [str(image_file)]
