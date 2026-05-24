import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.oxml import OxmlElement

from converter.formula_stats import inspect_document_formulas
from converter.models import ConvertResult


def test_inspect_document_formulas_counts_omml_and_residual_latex():
    doc = Document()
    doc.add_paragraph("残留 $x^2$")
    paragraph = doc.add_paragraph()
    paragraph._p.append(OxmlElement("m:oMath"))

    stats = inspect_document_formulas(doc)

    assert stats.document_total == 1
    assert stats.native_omml == 1
    assert stats.postprocessed == 0
    assert stats.residual_latex == 1


def test_inspect_document_formulas_subtracts_postprocessed_count():
    doc = Document()
    for _ in range(2):
        paragraph = doc.add_paragraph()
        paragraph._p.append(OxmlElement("m:oMath"))
    result = ConvertResult(converted=1)

    stats = inspect_document_formulas(doc, result)

    assert stats.document_total == 2
    assert stats.native_omml == 1
    assert stats.postprocessed == 1
