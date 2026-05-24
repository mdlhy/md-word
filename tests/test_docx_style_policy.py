from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from converter.pandoc_driver import is_pandoc_available
from converter.markdown_pipeline import convert_markdown_text_to_docx


def _rpr_from_style(doc, style_name):
    return doc.styles[style_name].element.get_or_add_rPr()


def _assert_songti_times_black(rpr):
    rfonts = rpr.find(qn("w:rFonts"))
    assert rfonts is not None
    assert rfonts.get(qn("w:eastAsia")) == "宋体"
    assert rfonts.get(qn("w:ascii")) == "Times New Roman"
    assert rfonts.get(qn("w:hAnsi")) == "Times New Roman"
    assert rfonts.get(qn("w:cs")) == "Times New Roman"
    assert rfonts.get(qn("w:eastAsiaTheme")) is None
    assert rfonts.get(qn("w:asciiTheme")) is None
    assert rfonts.get(qn("w:hAnsiTheme")) is None

    color = rpr.find(qn("w:color"))
    assert color is not None
    assert color.get(qn("w:val")) == "000000"
    assert color.get(qn("w:themeColor")) is None


def _run_rpr(run):
    rpr = run._element.find(qn("w:rPr"))
    assert rpr is not None
    return rpr


def test_reference_docs_use_songti_times_black_for_headings():
    for path in Path("templates").glob("*.docx"):
        doc = Document(str(path))
        for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle"):
            _assert_songti_times_black(_rpr_from_style(doc, style_name))


@pytest.mark.skipif(not is_pandoc_available(), reason="pandoc not installed")
def test_pandoc_output_subheading_is_not_blue(tmp_path):
    output_path = tmp_path / "font-policy.docx"
    md_text = (
        "## 问题 1：AFDM 融合公式缺失\n\n"
        "来源：专利审查意见 位置：3.2.3 节 描述：Softmax 归一化。\n"
    )

    convert_markdown_text_to_docx(
        md_text,
        str(output_path),
        template_name="academic",
        engine="pandoc",
    )

    doc = Document(str(output_path))
    heading = next(p for p in doc.paragraphs if p.style.name == "Heading 2")
    _assert_songti_times_black(_rpr_from_style(doc, "Heading 2"))
    for run in heading.runs:
        if run.text.strip():
            _assert_songti_times_black(_run_rpr(run))
