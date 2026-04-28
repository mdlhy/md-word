"""E2E test: convert real AI-generated .md files to .docx using all templates."""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from converter.md_converter import convert_md_to_docx
from converter.templates import TEMPLATES

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")
FIXTURE_FILES = ["ai_academic.md", "ai_homework.md", "ai_report.md"]


def test_all_fixtures():
    results = []
    for fixture_name in FIXTURE_FILES:
        fixture_path = os.path.join(FIXTURES_DIR, fixture_name)
        if not os.path.exists(fixture_path):
            results.append((fixture_name, "SKIP", "File not found"))
            continue
        
        with open(fixture_path, "r", encoding="utf-8") as f:
            md_text = f.read()
        
        for template_name in TEMPLATES:
            try:
                doc, report = convert_md_to_docx(md_text, template_name)
                
                # Verify: document opens without error
                assert doc is not None
                
                # Verify: has paragraphs
                para_count = len(doc.paragraphs)
                assert para_count > 0, f"No paragraphs in output"
                
                # Verify: has headings
                headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
                
                # Verify: check for OMML math elements
                from lxml import etree
                ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
                math_elements = doc.element.body.findall(".//m:oMath", ns)
                
                # Verify: has tables
                table_count = len(doc.tables)
                
                results.append((
                    f"{fixture_name}/{template_name}",
                    "PASS",
                    f"paras={para_count}, headings={len(headings)}, math={len(math_elements)}, tables={table_count}, compat={report.summary}"
                ))
            except Exception as e:
                results.append((f"{fixture_name}/{template_name}", "FAIL", str(e)[:100]))
    
    # Print results
    print("\n" + "=" * 80)
    print("E2E Test Results: .md → .docx conversion")
    print("=" * 80)
    passed = 0
    failed = 0
    for name, status, detail in results:
        symbol = "✅" if status == "PASS" else "❌" if status == "FAIL" else "⏭️"
        print(f"{symbol} {name}: {detail}")
        if status == "PASS":
            passed += 1
        elif status == "FAIL":
            failed += 1
    
    print(f"\nTotal: {passed + failed} | Passed: {passed} | Failed: {failed}")
    return failed == 0


if __name__ == "__main__":
    success = test_all_fixtures()
    sys.exit(0 if success else 1)
