"""Generate intentionally broken .docx files for repair testing."""
import os
from docx import Document
from docx.shared import Pt

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")
os.makedirs(FIXTURES_DIR, exist_ok=True)


def generate_broken_headings():
    """Create a .docx with # prefixed headings as plain text."""
    doc = Document()
    doc.add_paragraph("# 第一章 引言")
    doc.add_paragraph("这是引言部分的内容。")
    doc.add_paragraph("## 1.1 背景")
    doc.add_paragraph("研究背景如下。")
    doc.add_paragraph("### 1.1.1 详细背景")
    doc.add_paragraph("详细内容。")
    path = os.path.join(FIXTURES_DIR, "broken_headings.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


def generate_broken_lists():
    """Create a .docx with list-like plain text."""
    doc = Document()
    doc.add_paragraph("1. 第一项")
    doc.add_paragraph("2. 第二项")
    doc.add_paragraph("3. 第三项")
    doc.add_paragraph("- 无序项目A")
    doc.add_paragraph("- 无序项目B")
    doc.add_paragraph("正文段落")
    path = os.path.join(FIXTURES_DIR, "broken_lists.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


def generate_broken_mixed():
    """Create a .docx with multiple issues."""
    doc = Document()
    doc.add_paragraph("# 混合问题文档")
    doc.add_paragraph("1. 列表项一")
    doc.add_paragraph("2. 列表项二")
    doc.add_paragraph("正文内容 $x^2$ 公式。")
    doc.add_paragraph("## 第二节")
    doc.add_paragraph("- 项目A")
    doc.add_paragraph("- 项目B")
    path = os.path.join(FIXTURES_DIR, "broken_mixed.docx")
    doc.save(path)
    print(f"Created: {path}")
    return path


if __name__ == "__main__":
    generate_broken_headings()
    generate_broken_lists()
    generate_broken_mixed()
    print("All broken .docx files generated.")
