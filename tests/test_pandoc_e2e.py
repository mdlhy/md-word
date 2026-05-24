import os
import sys
from zipfile import ZipFile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest
from docx import Document

from converter.pandoc_driver import is_pandoc_available
from converter.markdown_pipeline import convert_markdown_text_to_docx
from converter.diagram_renderer import HAS_MATPLOTLIB


pytestmark = pytest.mark.skipif(
    not is_pandoc_available(),
    reason="pandoc not installed",
)


def test_pandoc_pipeline_generates_openable_docx(tmp_path):
    output_path = tmp_path / "pandoc.docx"
    md_text = (
        "# Title\n\n"
        "Paragraph with $x^2$.\n\n"
        "| A | B |\n"
        "|---|---|\n"
        "| 1 | 2 |\n"
    )

    result = convert_markdown_text_to_docx(
        md_text,
        str(output_path),
        template_name="academic",
        engine="pandoc",
    )

    doc = Document(str(output_path))

    assert result.engine == "pandoc"
    assert output_path.exists()
    assert any(p.style.name.startswith("Heading") for p in doc.paragraphs)
    assert len(doc.tables) == 1

    with ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "Title" in document_xml


@pytest.mark.skipif(not HAS_MATPLOTLIB, reason="matplotlib not installed")
def test_pandoc_pipeline_renders_custom_diagram_blocks(tmp_path):
    output_path = tmp_path / "diagram.docx"
    md_text = (
        "# Diagram\n\n"
        "```matrix\n"
        "name: A\n"
        "1 2\n"
        "3 4\n"
        "caption: Matrix A\n"
        "```\n"
    )

    result = convert_markdown_text_to_docx(
        md_text,
        str(output_path),
        template_name="academic",
        engine="pandoc",
    )
    doc = Document(str(output_path))
    text = "\n".join(p.text for p in doc.paragraphs)

    assert result.engine == "pandoc"
    assert len(doc.inline_shapes) >= 1
    assert "1 2" not in text
