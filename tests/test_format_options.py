import os
import sys

import pytest
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from converter.docx_repair import repair_docx
from converter.format_options import build_effective_template, list_format_presets, parse_format_options


def test_list_format_presets_exposes_common_choices():
    payload = list_format_presets()

    assert any(item["id"] == "academic" for item in payload["presets"])
    assert "宋体" in payload["choices"]["font_cn"]
    assert "1.5倍" in payload["choices"]["line_spacing"]


def test_build_effective_template_applies_user_options():
    config = build_effective_template(
        "academic",
        {
            "body": {
                "font_cn": "微软雅黑",
                "font_en": "Calibri",
                "size": "五号",
                "line_spacing": "1.25倍",
                "first_indent": "0字符",
            },
            "page": {"margin_preset": "narrow"},
            "heading": {"numbering": False, "alignment": "左对齐"},
            "table": {"three_line_default": False},
            "footer": {"page_number": False},
        },
    )

    assert config["body"]["font_cn"] == "微软雅黑"
    assert config["body"]["font_en"] == "Calibri"
    assert config["body"]["line_spacing"] == "1.25倍"
    assert config["page"]["margin_left"] == "1.8厘米"
    assert config["heading1"]["numbering"]["enabled"] is False
    assert config["table"]["three_line_default"] is False
    assert config["footer"]["page_number"] is False


def test_parse_format_options_rejects_invalid_json():
    with pytest.raises(ValueError):
        parse_format_options("{bad json")


def test_repair_docx_applies_custom_body_and_page(tmp_path):
    input_path = tmp_path / "input.docx"
    doc = Document()
    doc.add_paragraph("hello world")
    doc.save(input_path)

    repaired, _ = repair_docx(
        str(input_path),
        "academic",
        fix_formulas=False,
        format_options={
            "body": {
                "font_cn": "微软雅黑",
                "font_en": "Calibri",
                "size": "五号",
                "line_spacing": "1.25倍",
                "first_indent": "0字符",
            },
            "page": {"margin_preset": "narrow"},
        },
    )

    assert round(repaired.sections[0].left_margin.cm, 1) == 1.8
    run = repaired.paragraphs[0].runs[0]
    assert run.font.name == "Calibri"
    assert run.font.size.pt == pytest.approx(10.5)
