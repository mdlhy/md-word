"""MD → DOCX orchestrator.

Ties together markdown parsing, template loading, element converters,
and math OMML generation to produce a styled .docx with compatibility report.
"""

import copy
import logging
import os
from urllib.parse import urlparse

import math2docx
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from converter.md_parser import Token, parse_markdown
from converter.format_options import build_effective_template
from converter.compat_report import CompatItem, CompatReport, assess_formula_risk, assess_image_risk, generate_compat_report
from converter.elements.heading import add_heading
from converter.elements.table import add_table
from converter.elements.list import add_list_item
from converter.elements.blockquote import add_blockquote
from converter.elements.code import add_code_block
from converter.elements.image import add_image, add_image_bytes
from converter.format_units import to_pt
from converter.inline_text import tokens_to_plain_text
from converter.numbering import process_heading_numbering
from converter.diagram_renderer import is_diagram_language, render_diagram
from converter.elements.special_sections import (
    is_abstract_heading, is_references_heading, is_keywords_paragraph,
    format_abstract_heading, format_abstract_content, format_keywords_paragraph,
    format_references_heading, format_reference_entry,
)
from converter.elements.caption import add_figure_caption, add_table_caption

__all__ = ["convert_md_to_docx"]

logger = logging.getLogger(__name__)

NS_M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def _clear_document_content(doc: Document):
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}p") or child.tag.endswith("}tbl"):
            body.remove(child)
    doc.add_paragraph()


def _add_math_to_paragraph(paragraph, latex: str, display: bool):
    try:
        tmp_doc = Document()
        tmp_p = tmp_doc.add_paragraph()
        math2docx.add_math(tmp_p, latex)

        omath_elements = tmp_p._p.findall(f"{NS_M}oMath")
        if not omath_elements:
            raise ValueError("math2docx produced no oMath element")

        if display:
            omath_para = OxmlElement("m:oMathPara")
            for elem in omath_elements:
                omath_para.append(copy.deepcopy(elem))
            paragraph._p.append(omath_para)
        else:
            for elem in omath_elements:
                paragraph._p.append(copy.deepcopy(elem))
    except Exception:
        try:
            from converter.math_renderer import latex_to_omml, latex_to_omml_para

            fallback = latex_to_omml_para(latex) if display else latex_to_omml(latex)
            if fallback is None:
                raise ValueError("math_renderer produced no OMML")
            paragraph._p.append(copy.deepcopy(fallback))
            return
        except Exception:
            pass

        logger.warning(f"Math conversion failed for: {latex[:50]}")
        safe_text = latex.encode('utf-8', errors='replace').decode('utf-8')
        safe_text = ''.join(c for c in safe_text if ord(c) >= 32 or c in '\t\n')
        run = paragraph.add_run(safe_text)
        run.font.italic = True


def _clean_text(text: str) -> str:
    return ''.join(c for c in text if ord(c) >= 32 or c in '\t\n')


def _add_inline_children(paragraph, children: list[Token], config: dict):
    for child in children:
        if child.type == "text":
            paragraph.add_run(_clean_text(child.content))
        elif child.type == "strong":
            text = child.content or tokens_to_plain_text(child.children)
            run = paragraph.add_run(_clean_text(text))
            run.font.bold = True
        elif child.type == "em":
            text = child.content or tokens_to_plain_text(child.children)
            run = paragraph.add_run(_clean_text(text))
            run.font.italic = True
        elif child.type == "link":
            text = child.content or tokens_to_plain_text(child.children)
            paragraph.add_run(_clean_text(text))
        elif child.type == "codespan":
            run = paragraph.add_run(_clean_text(child.content))
            run.font.name = "Consolas"
            code_config = config.get("code", {})
            run.font.size = to_pt(code_config.get("size", "五号"))
        elif child.type == "math":
            _add_math_to_paragraph(
                paragraph, child.content, child.attrs.get("display", False)
            )
        elif child.type == "image":
            add_image(paragraph.part.document, child, config)
            alt_text = child.attrs.get("alt", "")
            if alt_text:
                add_figure_caption(paragraph.part.document, alt_text, config)
        elif child.children:
            _add_inline_children(paragraph, child.children, config)


def _add_rich_paragraph(doc: Document, token: Token, config: dict):
    p = doc.add_paragraph()
    _add_inline_children(p, token.children, config)
    _apply_body_format(p, config)
    return p


def _apply_body_format(paragraph, config: dict):
    body_cfg = config.get("body", {})
    if not body_cfg:
        return

    for run in paragraph.runs:
        if body_cfg.get("font_cn"):
            run.font.name = body_cfg.get("font_en", "Times New Roman")
            r_elem = run._element
            rpr = r_elem.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                r_elem.insert(0, rpr)
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            rfonts.set(qn("w:eastAsia"), body_cfg["font_cn"])
        if body_cfg.get("size"):
            run.font.size = to_pt(body_cfg["size"])

    if body_cfg.get("line_spacing"):
        from converter.format_units import to_spacing
        spacing = to_spacing(body_cfg["line_spacing"])
        paragraph.paragraph_format.line_spacing = spacing

    if body_cfg.get("first_indent"):
        from converter.format_units import to_length, font_size_to_pt
        body_size_pt = font_size_to_pt(body_cfg.get("size", "小四"))
        indent = to_length(body_cfg["first_indent"], font_size_pt=body_size_pt)
        paragraph.paragraph_format.first_line_indent = indent

    from converter.format_units import to_alignment
    alignment = to_alignment(body_cfg.get("alignment", "两端对齐"))
    if alignment is not None:
        paragraph.alignment = alignment


def _force_run_black(run):
    from docx.oxml.ns import qn as _qn
    rpr = run._element.find(_qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.insert(0, rpr)
    color_el = rpr.find(_qn("w:color"))
    if color_el is None:
        color_el = OxmlElement("w:color")
        rpr.append(color_el)
    color_el.set(_qn("w:val"), "000000")
    color_el.set(_qn("w:themeColor"), "text1")


_TOC_HEADING_RE = __import__("re").compile(
    r"^(目录|Table\s+of\s+Contents|Contents|TOC)$", __import__("re").IGNORECASE
)


def _is_toc_heading(text: str) -> bool:
    return bool(_TOC_HEADING_RE.match(text.strip()))


def _add_toc_field(doc: Document):
    """Insert a TOC field paragraph that generates a table of contents in Word."""
    p = doc.add_paragraph()

    run_begin = p.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run_begin._element.append(fldChar_begin)

    run_instr = p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run_instr._element.append(instrText)

    run_sep = p.add_run()
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    run_sep._element.append(fldChar_sep)

    placeholder = p.add_run("请更新目录：右键 → 更新域")
    _force_run_black(placeholder)

    run_end = p.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run_end._element.append(fldChar_end)

    return p


def _add_thematic_break(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "single")
        el.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz", "6")
        el.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space", "1")
        el.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color", "CCCCCC")
        pBdr.append(el)
    pPr.append(pBdr)
    return p


def _collect_math_compat(token: Token, items: list[CompatItem]):
    if token.type == "math":
        risk = assess_formula_risk(token.content)
        items.append(CompatItem(
            element_type="formula", risk=risk,
            description=f"Formula: {token.content[:60]}",
        ))
    for child in token.children:
        _collect_math_compat(child, items)


def _is_remote_url(src: str) -> bool:
    parsed = urlparse(src)
    return parsed.scheme in ("http", "https")


def _resolve_relative_image_sources(tokens: list[Token], base_dir: str | None):
    if not base_dir:
        return

    base_dir = os.path.abspath(base_dir)

    def walk(token: Token):
        if token.type == "image":
            src = token.attrs.get("src", "") or token.content
            if src and not _is_remote_url(src) and not os.path.isabs(src):
                token.attrs["src"] = os.path.abspath(os.path.join(base_dir, src))
        for child in token.children:
            walk(child)

    for token in tokens:
        walk(token)


def _add_diagram_block(doc: Document, token: Token, config: dict) -> bool:
    language = token.attrs.get("language") or token.attrs.get("lang", "")
    if not is_diagram_language(language):
        return False

    image_data, caption = render_diagram(token.content, language)
    if not image_data:
        return False

    add_image_bytes(doc, image_data, config)
    if caption:
        add_figure_caption(doc, caption, config)
    return True


def _apply_header_footer(doc: Document, config: dict):
    """Apply header (thesis title) and footer (page number) to all sections."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    header_cfg = config.get("header", {})
    footer_cfg = config.get("footer", {})

    first_h1_text = ""
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            first_h1_text = p.text.strip()
            break

    for section in doc.sections:
        if header_cfg.get("text") or first_h1_text:
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hp.text = header_cfg.get("text") or first_h1_text
            header_size = to_pt(header_cfg.get("size", "小五"))
            for run in hp.runs:
                run.font.size = header_size
                run.font.name = "Times New Roman"
                r_elem = run._element
                rpr = r_elem.find(qn("w:rPr"))
                if rpr is None:
                    rpr = OxmlElement("w:rPr")
                    r_elem.insert(0, rpr)
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.insert(0, rfonts)
                rfonts.set(qn("w:eastAsia"), header_cfg.get("font_cn", "宋体"))

        if footer_cfg.get("page_number", True):
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            footer_size = to_pt(footer_cfg.get("size", "小五"))
            run = fp.add_run()
            run.font.size = footer_size

            fldChar_begin = OxmlElement("w:fldChar")
            fldChar_begin.set(qn("w:fldCharType"), "begin")
            run._element.append(fldChar_begin)

            instr_run = fp.add_run()
            instr_run.font.size = footer_size
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = " PAGE "
            instr_run._element.append(instrText)

            end_run = fp.add_run()
            end_run.font.size = footer_size
            fldChar_end = OxmlElement("w:fldChar")
            fldChar_end.set(qn("w:fldCharType"), "end")
            end_run._element.append(fldChar_end)


def _apply_title_author(doc: Document, config: dict, heading_paragraphs: list):
    """Apply title/author styles for templates that define them (e.g. dialectics).

    Detects the first H1 as title, the bold paragraph immediately after as author.
    Only applies if the template has 'title' and/or 'author' config blocks.
    """
    title_cfg = config.get("title")
    author_cfg = config.get("author")
    if not title_cfg and not author_cfg:
        return

    if not heading_paragraphs:
        return

    first_level, first_para = heading_paragraphs[0]
    if first_level != 1:
        return

    if title_cfg:
        for run in first_para.runs:
            if title_cfg.get("font_cn"):
                run.font.name = title_cfg.get("font_en", "Times New Roman")
                r_elem = run._element
                rpr = r_elem.find(qn("w:rPr"))
                if rpr is None:
                    rpr = OxmlElement("w:rPr")
                    r_elem.insert(0, rpr)
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.insert(0, rfonts)
                rfonts.set(qn("w:eastAsia"), title_cfg["font_cn"])
            if title_cfg.get("size"):
                run.font.size = to_pt(title_cfg["size"])
            run.font.bold = False
            _force_run_black(run)
        from converter.format_units import to_alignment
        alignment = to_alignment(title_cfg.get("alignment", "居中对齐"))
        if alignment is not None:
            first_para.alignment = alignment

    if author_cfg:
        body = doc.element.body
        first_para_elem = first_para._element
        next_elem = first_para_elem.getnext()
        if next_elem is not None:
            for para in doc.paragraphs:
                if para._element is next_elem and para.runs:
                    all_bold = all(r.font.bold for r in para.runs if r.text.strip())
                    if all_bold:
                        for run in para.runs:
                            if author_cfg.get("font_cn"):
                                run.font.name = author_cfg.get("font_en", "Times New Roman")
                                r_elem = run._element
                                rpr = r_elem.find(qn("w:rPr"))
                                if rpr is None:
                                    rpr = OxmlElement("w:rPr")
                                    r_elem.insert(0, rpr)
                                rfonts = rpr.find(qn("w:rFonts"))
                                if rfonts is None:
                                    rfonts = OxmlElement("w:rFonts")
                                    rpr.insert(0, rfonts)
                                rfonts.set(qn("w:eastAsia"), author_cfg["font_cn"])
                            if author_cfg.get("size"):
                                run.font.size = to_pt(author_cfg["size"])
                            run.font.bold = False
                            _force_run_black(run)
                        from converter.format_units import to_alignment
                        alignment = to_alignment(author_cfg.get("alignment", "居中对齐"))
                        if alignment is not None:
                            para.alignment = alignment
                    break


def convert_md_to_docx(
    md_text: str,
    template_name: str = "academic",
    three_line: bool = False,
    base_dir: str | None = None,
    format_options: dict | None = None,
) -> tuple[Document, CompatReport]:
    """Convert markdown text to a python-docx Document with compatibility report.

    Args:
        md_text: Markdown source text.
        template_name: Template name (academic, homework, report).
        three_line: Force three-line table style.
        base_dir: Base directory for resolving relative image paths.

    Returns:
        (Document, CompatReport) tuple.
    """
    md_text = _clean_text(md_text)
    tokens = parse_markdown(md_text)
    _resolve_relative_image_sources(tokens, base_dir)

    template_path = os.path.join(
        os.path.dirname(os.path.dirname(__file__)), "templates", f"{template_name}.docx"
    )
    doc = Document(template_path)
    _clear_document_content(doc)

    config = build_effective_template(
        template_name,
        format_options,
        three_line_override=three_line if three_line else None,
    )
    compat_items: list[CompatItem] = []
    heading_paragraphs: list[tuple[int, object]] = []
    section_state = {"in_abstract": False, "in_references": False}

    for token in tokens:
        if token.type == "heading":
            heading_text = token.content.strip()
            if not heading_text and token.children:
                heading_text = tokens_to_plain_text(token.children).strip()

            if _is_toc_heading(heading_text):
                section_state["in_abstract"] = False
                section_state["in_references"] = False
                _add_toc_field(doc)
                continue

            p = add_heading(doc, token, config)
            level = token.attrs.get("level", token.level) or 1
            if level > 3:
                level = 3

            if is_abstract_heading(heading_text):
                section_state["in_abstract"] = True
                section_state["in_references"] = False
                format_abstract_heading(p, config)
            elif is_references_heading(heading_text):
                section_state["in_references"] = True
                section_state["in_abstract"] = False
                format_references_heading(p, config)
            else:
                section_state["in_abstract"] = False
                section_state["in_references"] = False
                heading_paragraphs.append((level, p))

        elif token.type == "table":
            add_table(doc, token, config, three_line)
            compat_items.append(CompatItem(
                element_type="table", risk="low",
                description="Table element (low WPS risk)",
            ))

        elif token.type == "list":
            add_list_item(doc, token, config)

        elif token.type == "blockquote":
            add_blockquote(doc, token, config)

        elif token.type == "code":
            if not _add_diagram_block(doc, token, config):
                add_code_block(doc, token, config)

        elif token.type == "image":
            add_image(doc, token, config)
            alt_text = token.attrs.get("alt", "") or token.content
            if alt_text:
                add_figure_caption(doc, alt_text, config)
            src = token.attrs.get("src", "")
            risk = assess_image_risk(src)
            compat_items.append(CompatItem(
                element_type="image", risk=risk,
                description=f"Image: {src}",
            ))

        elif token.type == "math":
            _add_math_to_paragraph(
                doc.add_paragraph(), token.content, token.attrs.get("display", False)
            )
            risk = assess_formula_risk(token.content)
            compat_items.append(CompatItem(
                element_type="formula", risk=risk,
                description=f"Formula: {token.content[:60]}",
            ))

        elif token.type == "paragraph":
            _collect_math_compat(token, compat_items)
            if token.children:
                p = _add_rich_paragraph(doc, token, config)
            else:
                p = doc.add_paragraph(token.content)
                _apply_body_format(p, config)

            if section_state["in_abstract"]:
                para_text = p.text.strip()
                if is_keywords_paragraph(para_text):
                    format_keywords_paragraph(p, config)
                else:
                    format_abstract_content(p, config)
            elif section_state["in_references"]:
                format_reference_entry(p, config)

        elif token.type == "thematic_break":
            _add_thematic_break(doc)

    process_heading_numbering(doc, config, heading_paragraphs)
    _apply_title_author(doc, config, heading_paragraphs)
    _apply_header_footer(doc, config)

    report = generate_compat_report(compat_items)
    return doc, report


def convert_directory(input_dir, output_dir=None, template_name="academic",
                      three_line=False, recursive=False):
    """Batch convert all .md files in a directory to .docx.

    Args:
        input_dir: Directory containing .md files
        output_dir: Output directory (defaults to input_dir)
        template_name: Template to use for conversion
        three_line: Use three-line table style
        recursive: Search subdirectories recursively

    Returns:
        List of (input_path, output_path, success, error) tuples
    """
    import glob as glob_module

    input_dir = os.path.abspath(input_dir)
    if not os.path.isdir(input_dir):
        raise NotADirectoryError(f"Not a directory: {input_dir}")

    if output_dir is None:
        output_dir = input_dir
    output_dir = os.path.abspath(output_dir)
    os.makedirs(output_dir, exist_ok=True)

    pattern = os.path.join(input_dir, "**", "*.md") if recursive else os.path.join(input_dir, "*.md")
    md_files = sorted(glob_module.glob(pattern, recursive=recursive))

    if not md_files:
        logger.warning(f"No .md files found in: {input_dir}")
        return []

    logger.info(f"Found {len(md_files)} markdown file(s)")

    results = []
    for md_file in md_files:
        rel_path = os.path.relpath(md_file, input_dir)
        out_path = os.path.join(output_dir, os.path.splitext(rel_path)[0] + ".docx")
        os.makedirs(os.path.dirname(out_path), exist_ok=True)

        try:
            with open(md_file, "r", encoding="utf-8") as f:
                md_text = f.read()
            doc, _ = convert_md_to_docx(
                md_text,
                template_name,
                three_line,
                base_dir=os.path.dirname(md_file),
            )
            doc.save(out_path)
            results.append((md_file, out_path, True, None))
            logger.info(f"Converted: {md_file} → {out_path}")
        except Exception as e:
            logger.error(f"Failed to convert {md_file}: {e}")
            results.append((md_file, out_path, False, str(e)))

    success = sum(1 for *_, ok, _ in results if ok)
    failed = len(results) - success
    logger.info(f"Conversion complete: {success} succeeded, {failed} failed")

    return results
