import os
import tempfile
from io import BytesIO

import requests
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from converter.md_parser import Token
from converter.format_units import parse_unit


def add_image(doc: Document, token: Token, template_config: dict):
    src = token.attrs.get("src", "")
    if not src:
        src = token.content
    
    if not src:
        return None
    
    page_config = template_config.get("page", {})
    margin_left = parse_unit(page_config.get("margin_left", "3.17厘米"))[0]
    margin_right = parse_unit(page_config.get("margin_right", "3.17厘米"))[0]
    page_width = 21.0
    max_width_cm = page_width - margin_left - margin_right - 2.0
    max_width = Cm(max_width_cm)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    image_data = None
    temp_path = None
    
    try:
        if src.startswith(("http://", "https://")):
            try:
                resp = requests.get(src, timeout=10, stream=True)
                resp.raise_for_status()
                image_data = resp.content
            except Exception:
                run = p.add_run(f"图片加载失败: {src}")
                run.font.size = Pt(10)
                return p
        elif os.path.isfile(src):
            with open(src, "rb") as f:
                image_data = f.read()
        else:
            run = p.add_run(f"图片未找到: {src}")
            run.font.size = Pt(10)
            return p
        
        src_lower = src.lower()
        
        if src_lower.endswith(".svg"):
            # SVG → PNG via cairosvg, fallback to svglib
            png_data = _convert_svg_to_png(image_data)
            if png_data is None:
                run = p.add_run(f"SVG转换失败: {src}")
                run.font.size = Pt(10)
                return p
            image_data = png_data
        
        elif src_lower.endswith(".webp"):
            # WebP → PNG via Pillow
            png_data = _convert_webp_to_png(image_data)
            if png_data is None:
                run = p.add_run(f"WebP转换失败: {src}")
                run.font.size = Pt(10)
                return p
            image_data = png_data
        
        stream = BytesIO(image_data)
        run = p.add_run()
        run.add_picture(stream, width=max_width)
        return p
        
    except Exception:
        run = p.add_run(f"图片加载失败: {src}")
        run.font.size = Pt(10)
        return p


def _convert_svg_to_png(svg_data: bytes) -> bytes | None:
    """Convert SVG to PNG via cairosvg, falling back to svglib+reportlab."""
    try:
        import cairosvg
        return cairosvg.svg2png(bytestring=svg_data, output_width=800)
    except Exception:
        pass
    
    try:
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        
        with tempfile.NamedTemporaryFile(suffix=".svg", delete=False) as tmp:
            tmp.write(svg_data)
            tmp_path = tmp.name
        
        try:
            drawing = svg2rlg(tmp_path)
            if drawing is not None:
                png_data = renderPM.drawToString(drawing, fmt="PNG", dpi=150)
                return png_data
        finally:
            os.unlink(tmp_path)
    except Exception:
        pass
    
    return None


def _convert_webp_to_png(webp_data: bytes) -> bytes | None:
    try:
        from PIL import Image
        img = Image.open(BytesIO(webp_data))
        if img.mode == "RGBA":
            img = img.convert("RGB")
        buf = BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None
