from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

from converter.md_parser import Token
from converter.format_units import to_length, font_size_to_pt, to_pt


def _render_inline_children(paragraph, children, config):
    for child in children:
        if child.type == "text":
            paragraph.add_run(child.content)
        elif child.type == "strong":
            text = child.content or "".join(
                c.content for c in child.children if c.type == "text"
            )
            run = paragraph.add_run(text)
            run.font.bold = True
        elif child.type == "em":
            text = child.content or "".join(
                c.content for c in child.children if c.type == "text"
            )
            run = paragraph.add_run(text)
            run.font.italic = True
        elif child.type == "codespan":
            run = paragraph.add_run(child.content)
            run.font.name = "Consolas"
            code_config = config.get("code", {})
            run.font.size = to_pt(code_config.get("size", "五号"))
        elif child.type == "math":
            run = paragraph.add_run(child.content)
            run.font.italic = True


def add_blockquote(doc: Document, token: Token, template_config: dict):
    """Add a blockquote from a parsed markdown token to a docx document.
    
    - Left indent: template config quote.indent (in cm, default 2)
    - Left border: template config quote.border_color (default CCCCCC), width 3pt
    - Each line is an independent paragraph sharing indent and border
    """
    quote_config = template_config.get("quote", {})
    indent_value = quote_config.get("indent", "2字符")
    border_color = quote_config.get("border_color", "CCCCCC")
    
    body_config = template_config.get("body", {})
    body_size_pt = font_size_to_pt(body_config.get("size", "小四"))
    
    paragraphs = []
    
    def add_quote_paragraph_from_token(child_token):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = to_length(indent_value, font_size_pt=body_size_pt)
        p.paragraph_format.space_after = Pt(4)
        
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left_border = OxmlElement("w:left")
        left_border.set(qn("w:val"), "single")
        left_border.set(qn("w:sz"), "24")
        left_border.set(qn("w:space"), "4")
        left_border.set(qn("w:color"), border_color)
        pBdr.append(left_border)
        pPr.append(pBdr)
        
        _render_inline_children(p, child_token.children, template_config)
        
        paragraphs.append(p)
        return p
    
    for child in token.children:
        if child.type == "paragraph":
            add_quote_paragraph_from_token(child)
        elif child.type == "text":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = to_length(indent_value, font_size_pt=body_size_pt)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left_border = OxmlElement("w:left")
            left_border.set(qn("w:val"), "single")
            left_border.set(qn("w:sz"), "24")
            left_border.set(qn("w:space"), "4")
            left_border.set(qn("w:color"), border_color)
            pBdr.append(left_border)
            pPr.append(pBdr)
            run = p.add_run(child.content)
            paragraphs.append(p)
        elif child.type == "blockquote":
            nested_indent_cm = to_length(indent_value, font_size_pt=body_size_pt).cm + 2
            nested_config = dict(template_config)
            nested_config["quote"] = dict(quote_config, indent=f"{nested_indent_cm}厘米")
            nested_ps = add_blockquote(doc, child, nested_config)
            paragraphs.extend(nested_ps)
    
    if not paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = to_length(indent_value, font_size_pt=body_size_pt)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left_border = OxmlElement("w:left")
        left_border.set(qn("w:val"), "single")
        left_border.set(qn("w:sz"), "24")
        left_border.set(qn("w:space"), "4")
        left_border.set(qn("w:color"), border_color)
        pBdr.append(left_border)
        pPr.append(pBdr)
        run = p.add_run(token.content)
        paragraphs.append(p)
    
    return paragraphs
