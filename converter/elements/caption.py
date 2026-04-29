"""Figure and table caption auto-numbering via Word SEQ fields.

Inserts '图 1 Caption text' or '表 1 Caption text' paragraphs using
SEQ field codes so numbering updates automatically in Word/WPS.

Usage:
    from converter.elements.caption import add_figure_caption, add_table_caption
    add_figure_caption(doc, "实验流程图", config)
    add_table_caption(doc, "实验结果", config)
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from converter.format_units import to_pt

_FIGURE_SEQ = "Figure"
_TABLE_SEQ = "Table"


def _add_seq_field(paragraph, seq_name: str):
    """Insert a SEQ field into a paragraph run.

    Generates XML equivalent to: {SEQ seq_name \\* ARABIC}
    """
    run = paragraph.add_run()

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar_begin)

    instrText_run = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" SEQ {seq_name} \\* ARABIC "
    instrText_run.append(instrText)
    paragraph._p.append(instrText_run)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    end_run = OxmlElement("w:r")
    end_run.append(fldChar_end)
    paragraph._p.append(end_run)


def add_figure_caption(doc: Document, caption_text: str, config: dict):
    """Add a figure caption paragraph below an image.

    Format: '图 {SEQ Figure} caption_text' — centered, 五号 font.
    """
    fig_cfg = config.get("figure", {})
    prefix = fig_cfg.get("caption_prefix", "图")
    caption_size = fig_cfg.get("caption_size", "五号")
    position = fig_cfg.get("caption_position", "below")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    label_run = p.add_run(f"{prefix} ")
    label_run.font.size = to_pt(caption_size)

    _add_seq_field(p, _FIGURE_SEQ)

    if caption_text:
        sep_run = p.add_run(f" {caption_text}")
        sep_run.font.size = to_pt(caption_size)

    return p


def add_table_caption(doc: Document, caption_text: str, config: dict):
    """Add a table caption paragraph above a table.

    Format: '表 {SEQ Table} caption_text' — centered, 五号 font.
    """
    tbl_cfg = config.get("table", {})
    prefix = tbl_cfg.get("caption_prefix", "表")
    caption_size = tbl_cfg.get("caption_size", "五号")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    label_run = p.add_run(f"{prefix} ")
    label_run.font.size = to_pt(caption_size)

    _add_seq_field(p, _TABLE_SEQ)

    if caption_text:
        sep_run = p.add_run(f" {caption_text}")
        sep_run.font.size = to_pt(caption_size)

    return p
