from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from converter.md_parser import Token
from converter.format_units import to_pt
from converter.style_presets import (
    CODE_FONT, CODE_FONT_SIZE, CODE_BG_COLOR,
    CODE_LINE_SPACING, CODE_INDENT_LEFT, CODE_INDENT_RIGHT,
    CODE_SPACE_BEFORE, CODE_SPACE_AFTER,
    SYNTAX_COLORS, LANGUAGE_ALIASES, get_syntax_color,
)

try:
    from pygments import lex
    from pygments.lexers import get_lexer_by_name, TextLexer
    HAS_PYGMENTS = True
except ImportError:
    HAS_PYGMENTS = False


def _get_lexer(language):
    lang = language.strip().lower() if language else ""
    lang = LANGUAGE_ALIASES.get(lang, lang)
    try:
        return get_lexer_by_name(lang, stripall=False)
    except Exception:
        try:
            return get_lexer_by_name("text", stripall=False)
        except Exception:
            return TextLexer()


def _tokenize_code(code, language):
    if not HAS_PYGMENTS:
        return [(None, line) for line in code.split("\n")]
    lexer = _get_lexer(language)
    return list(lex(code, lexer))


def _add_run_with_color(para, text, font_name, font_size, color=None):
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    if color:
        run.font.color.rgb = color
    rPr = run._element.find(qn("w:rPr"))
    if rPr is not None:
        rfonts = rPr.find(qn("w:rFonts"))
        if rfonts is not None:
            rfonts.set(qn("w:eastAsia"), font_name)
    return run


def add_code_block(doc: Document, token: Token, template_config: dict):
    code_config = template_config.get("code", {})
    font_name = code_config.get("font", CODE_FONT)
    font_size = to_pt(code_config.get("size", 10)) if code_config.get("size") else CODE_FONT_SIZE
    bg_color = code_config.get("bg_color", CODE_BG_COLOR)

    raw_code = token.content
    if not raw_code:
        for child in token.children:
            if child.type == "text":
                raw_code += child.content
    if not raw_code:
        return []

    language = token.attrs.get("language") or token.attrs.get("lang", "")
    tokens = _tokenize_code(raw_code, language)
    paragraphs = []

    if HAS_PYGMENTS and language:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = CODE_INDENT_LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = CODE_LINE_SPACING
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_color)
        pPr.append(shd)
        label_run = p.add_run(f"  {language}")
        label_run.font.name = font_name
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        label_run.font.italic = True
        paragraphs.append(p)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = CODE_INDENT_LEFT
    p.paragraph_format.right_indent = CODE_INDENT_RIGHT
    p.paragraph_format.space_before = CODE_SPACE_BEFORE
    p.paragraph_format.space_after = CODE_SPACE_AFTER
    p.paragraph_format.line_spacing = CODE_LINE_SPACING
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg_color)
    pPr.append(shd)

    if HAS_PYGMENTS and language:
        for token_type, token_value in tokens:
            if not token_value:
                continue
            color = get_syntax_color(token_type)
            parts = token_value.split("\n")
            for idx, part in enumerate(parts):
                if part:
                    _add_run_with_color(p, part, font_name, font_size, color)
                if idx < len(parts) - 1:
                    br_run = p.add_run()
                    br_run.add_break()
    else:
        lines = raw_code.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                br_run = p.add_run()
                br_run.add_break()
            if line:
                _add_run_with_color(p, line, font_name, font_size)

    paragraphs.append(p)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    paragraphs.append(spacer)

    return paragraphs
