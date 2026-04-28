from docx import Document
from docx.shared import Cm, Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from converter.md_parser import Token


def _set_cell_border(cell, **kwargs):
    """Set cell border. kwargs keys: top, bottom, left, right, insideH, insideV.
    Each value is a dict: {sz, color, val} or None to remove."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, props in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        if props is None:
            tcBorders.remove(element)
            continue
        for attr, val in props.items():
            element.set(qn(f"w:{attr}"), str(val))


def _set_table_borders(table, three_line=False):
    """Set table borders. three_line=True for academic style (top/bottom thick, header-bottom thin, no vertical)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    
    if three_line:
        # Top border: thick
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "12")
        top.set(qn("w:space"), "0")
        top.set(qn("w:color"), "000000")
        tblBorders.append(top)
        # Bottom border: thick
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")
        tblBorders.append(bottom)
        # No left/right/insideV
        for edge in ["left", "right", "insideV"]:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            tblBorders.append(el)
        # No insideH (data rows have no horizontal lines)
        insideH = OxmlElement("w:insideH")
        insideH.set(qn("w:val"), "nil")
        tblBorders.append(insideH)
        
        # Add thin border below header row (on the cells of the first row)
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                _set_cell_border(cell, bottom={"val": "single", "sz": "4", "space": "0", "color": "000000"})
    else:
        # Standard borders
        for edge in ["top", "bottom", "left", "right", "insideH", "insideV"]:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            tblBorders.append(el)


def _is_cell_token(token: Token) -> bool:
    return "head" in token.attrs


def _get_cell_text(token: Token) -> str:
    if token.content:
        return token.content
    if token.children:
        parts = []
        for child in token.children:
            if child.type == "text":
                parts.append(child.content)
            elif child.type == "math":
                parts.append(child.content)
            elif child.type in ("strong", "em", "codespan"):
                parts.append(child.content)
        return " ".join(parts).strip()
    return ""


def _is_row_of_cells(node: Token) -> bool:
    if not node.children:
        return False
    return any(_is_cell_token(c) for c in node.children)


def _extract_rows_v2(token: Token) -> tuple[list[list[str]], bool]:
    """Extract rows from a table Token.
    
    Handles mistune AST structure:
      table > table_head > table_cell (header cells, no table_row wrapper)
      table > table_body > table_row > table_cell
    """
    rows = []
    has_header = False

    def walk(node):
        nonlocal has_header
        if node.type == "table_head":
            has_header = True
            # table_head children are table_cell directly (no table_row wrapper)
            cells = []
            for child in node.children:
                if child.type == "table_cell":
                    cells.append(_get_cell_text(child))
                elif child.type == "text" and child.children:
                    for cc in child.children:
                        if cc.type == "table_cell":
                            cells.append(_get_cell_text(cc))
            if cells:
                rows.append(cells)
            return

        if node.type == "table_row":
            cells = []
            for child in node.children:
                if child.type == "table_cell":
                    cells.append(_get_cell_text(child))
            if cells:
                rows.append(cells)
            return

        for child in node.children:
            walk(child)

    walk(token)
    return rows, has_header


def add_table(doc: Document, token: Token, template_config: dict, three_line: bool = False):
    """Add a table from a parsed markdown token to a docx document.
    
    - Fixed column widths based on page size
    - Header row bold
    - Three-line table option (三线表)
    """
    rows_data, has_header = _extract_rows_v2(token)
    
    if not rows_data:
        return None
    
    num_cols = max(len(row) for row in rows_data)
    # Normalize rows to have same number of columns
    for row in rows_data:
        while len(row) < num_cols:
            row.append("")
    
    # Calculate column width
    page_config = template_config.get("page", {})
    margin_left = page_config.get("margin_left", 3.17)
    margin_right = page_config.get("margin_right", 3.17)
    page_width = 21.0  # A4 width in cm
    available_width = page_width - margin_left - margin_right
    col_width = Cm(available_width / num_cols)
    
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.autofit = False
    
    # Set column widths
    for col_idx in range(num_cols):
        for row in table.rows:
            row.cells[col_idx].width = col_width
    
    # Fill data
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            # Bold header row
            if row_idx == 0 and has_header:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    # Set borders
    table_config = template_config.get("table", {})
    use_three_line = three_line or table_config.get("three_line_default", False)
    _set_table_borders(table, three_line=use_three_line)
    
    return table
