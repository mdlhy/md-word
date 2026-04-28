"""Walk all paragraph containers in a .docx document."""
from docx import Document
from docx.text.paragraph import Paragraph
from typing import Iterator


def walk_all_paragraphs(doc: Document) -> Iterator[Paragraph]:
    """Yield all paragraphs from body, tables, headers, and footers."""
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        yield from _walk_table(table)

    for section in doc.sections:
        for attr in ['header', 'footer', 'first_page_header',
                     'first_page_footer', 'even_page_header', 'even_page_footer']:
            hf = getattr(section, attr, None)
            if hf and not hf.is_linked_to_previous:
                for paragraph in hf.paragraphs:
                    yield paragraph
                for table in hf.tables:
                    yield from _walk_table(table)


def _walk_table(table) -> Iterator[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
