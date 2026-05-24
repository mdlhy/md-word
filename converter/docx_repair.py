"""DOCX repair pipeline: detect and fix formatting issues in .docx files.

Applies template formatting (fonts, sizes, spacing, margins) plus heuristic
detection of structural issues (bare # headings, unstyled lists, etc.).
"""

import re
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

from converter.compat_report import CompatItem, CompatReport, generate_compat_report
from converter.format_options import build_effective_template
from converter.format_units import to_pt, to_length, to_spacing, font_size_to_pt, to_alignment, parse_unit
from converter.docx_style import apply_document_text_policy, set_run_text_policy


def repair_docx(
    input_path: str,
    template_name: str = "academic",
    fix_formulas: bool = True,
    format_options: dict | None = None,
    three_line_override: bool | None = None,
) -> tuple[Document, CompatReport]:
    doc = Document(input_path)
    template_config = build_effective_template(
        template_name,
        format_options,
        three_line_override=three_line_override,
    )
    items = []

    _apply_page_margins(doc, template_config)
    items.extend(_fix_headings(doc, template_config))
    items.extend(_fix_lists(doc, template_config))
    items.extend(_fix_tables(doc, template_config))
    _apply_body_formatting(doc, template_config)
    if fix_formulas:
        items.extend(_fix_formulas(doc))
    body_config = template_config.get("body", {})
    apply_document_text_policy(
        doc,
        cn_font=body_config.get("font_cn", "宋体"),
        en_font=body_config.get("font_en", "Times New Roman"),
    )

    report = generate_compat_report(items)
    return doc, report


# ---------------------------------------------------------------------------
# Page margins
# ---------------------------------------------------------------------------

def _apply_page_margins(doc: Document, template_config: dict):
    """Apply template page margins to all sections."""
    page = template_config.get("page", {})
    for section in doc.sections:
        if "margin_top" in page:
            section.top_margin = to_length(page["margin_top"])
        if "margin_bottom" in page:
            section.bottom_margin = to_length(page["margin_bottom"])
        if "margin_left" in page:
            section.left_margin = to_length(page["margin_left"])
        if "margin_right" in page:
            section.right_margin = to_length(page["margin_right"])


# ---------------------------------------------------------------------------
# Body formatting
# ---------------------------------------------------------------------------

def _apply_body_formatting(doc: Document, template_config: dict):
    """Apply template body font/size/spacing to all Normal paragraphs."""
    body_config = template_config.get("body", {})
    if not body_config:
        return

    font_cn = body_config.get("font_cn")
    font_en = body_config.get("font_en")
    font_size = body_config.get("size")
    font_size_pt = font_size_to_pt(font_size) if font_size else None
    line_spacing = body_config.get("line_spacing")
    first_indent = body_config.get("first_indent", "0字符")

    for para in doc.paragraphs:
        # Skip headings — they have their own formatting
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            continue

        for run in para.runs:
            if font_en:
                run.font.name = font_en
                # Set East Asian font via XML
                r_elem = run._element
                rpr = r_elem.find(qn("w:rPr"))
                if rpr is None:
                    rpr = OxmlElement("w:rPr")
                    r_elem.insert(0, rpr)
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.insert(0, rfonts)
                if font_cn:
                    rfonts.set(qn("w:eastAsia"), font_cn)

            if font_size:
                run.font.size = to_pt(font_size)

        # Paragraph-level formatting
        if line_spacing:
            pf = para.paragraph_format
            spacing_val = to_spacing(line_spacing)
            if isinstance(spacing_val, float):
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = spacing_val
            else:
                pf.line_spacing = spacing_val

        indent_num, indent_unit = parse_unit(first_indent)
        if indent_num > 0 and indent_unit == "char":
            para.paragraph_format.first_line_indent = to_length(first_indent, font_size_pt=font_size_pt or 12)
        elif indent_num > 0:
            para.paragraph_format.first_line_indent = to_length(first_indent)


# ---------------------------------------------------------------------------
# Heading fixes
# ---------------------------------------------------------------------------

def _fix_headings(doc, template_config) -> list[CompatItem]:
    items = []
    heading_re = re.compile(r'^(#{1,6})\s+(.+)$')

    for para in doc.paragraphs:
        if para.style.name != "Normal":
            continue

        text = para.text.strip()
        m = heading_re.match(text)
        if not m:
            continue

        level = min(len(m.group(1)), 3)
        title_text = m.group(2).strip()

        try:
            style_name = f"Heading {level}"
            _ = doc.styles[style_name]
            para.style = doc.styles[style_name]

            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = title_text
            else:
                para.add_run(title_text)

            # Apply template heading formatting
            heading_config = template_config.get(f"heading{level}", {})
            _apply_heading_formatting(para, heading_config)

            items.append(CompatItem(
                element_type="heading",
                risk="low",
                description=f"标题修复: '{text[:30]}' → {style_name}"
            ))
        except KeyError:
            pass

    # Also format existing headings with template styles
    for para in doc.paragraphs:
        if not (para.style and para.style.name and para.style.name.startswith("Heading")):
            continue
        # Extract level from style name
        try:
            level = int(para.style.name.split()[-1])
        except (ValueError, IndexError):
            continue
        if level > 3:
            continue
        heading_config = template_config.get(f"heading{level}", {})
        _apply_heading_formatting(para, heading_config)

    return items


def _apply_heading_formatting(para, heading_config: dict):
    """Apply template heading font/size/bold/center to a heading paragraph."""
    font_cn = heading_config.get("font_cn")
    font_en = heading_config.get("font_en")

    for run in para.runs:
        if font_en:
            set_run_text_policy(run, font_cn or "宋体", font_en)

        if heading_config.get("size"):
            run.font.size = to_pt(heading_config["size"])
        if heading_config.get("bold"):
            run.font.bold = True

    alignment = to_alignment(heading_config.get("alignment", "左对齐"))
    if alignment is not None:
        para.alignment = alignment


# ---------------------------------------------------------------------------
# List fixes
# ---------------------------------------------------------------------------

def _fix_lists(doc, template_config) -> list[CompatItem]:
    items = []
    ordered_re = re.compile(r'^(\d+)\.\s+(.+)$')
    unordered_re = re.compile(r'^[-*+]\s+(.+)$')

    for para in doc.paragraphs:
        if para.style.name not in ("Normal", "List Number", "List Bullet"):
            continue

        text = para.text.strip()
        m_ord = ordered_re.match(text)
        m_unord = unordered_re.match(text)

        if m_ord:
            item_text = m_ord.group(2)
            try:
                para.style = doc.styles["List Number"]
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = item_text
                else:
                    para.add_run(item_text)
                items.append(CompatItem(
                    element_type="list",
                    risk="low",
                    description=f"列表修复: '{text[:30]}' → List Number"
                ))
            except KeyError:
                pass
        elif m_unord:
            item_text = m_unord.group(1)
            try:
                para.style = doc.styles["List Bullet"]
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = item_text
                else:
                    para.add_run(item_text)
                items.append(CompatItem(
                    element_type="list",
                    risk="low",
                    description=f"列表修复: '{text[:30]}' → List Bullet"
                ))
            except KeyError:
                pass

    return items


# ---------------------------------------------------------------------------
# Table fixes
# ---------------------------------------------------------------------------

def _fix_tables(doc, template_config) -> list[CompatItem]:
    items = []

    page_config = template_config.get("page", {})
    margin_left = parse_unit(page_config.get("margin_left", "3.17厘米"))[0]
    margin_right = parse_unit(page_config.get("margin_right", "3.17厘米"))[0]
    available_width = 21.0 - margin_left - margin_right

    for table in doc.tables:
        num_cols = len(table.columns)
        if num_cols == 0:
            continue

        col_width = Cm(available_width / num_cols)
        table.autofit = False

        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width

        from converter.elements.table import _set_table_borders
        table_config = template_config.get("table", {})
        three_line = table_config.get("three_line_default", False)
        _set_table_borders(table, three_line=three_line)

        if table.rows:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

        items.append(CompatItem(
            element_type="table",
            risk="low",
            description="表格修复: 固定列宽 + 边框优化"
        ))

    return items


# ---------------------------------------------------------------------------
# Formula fixes
# ---------------------------------------------------------------------------

def _fix_formulas(doc) -> list[CompatItem]:
    items = []
    try:
        from converter.orchestrator import convert_docx_in_memory
        result = convert_docx_in_memory(doc)
        if result.converted > 0:
            items.append(CompatItem(
                element_type="formula",
                risk="low",
                description=f"公式修复: {result.converted} 个公式转换成功"
            ))
        if result.failed > 0:
            items.append(CompatItem(
                element_type="formula",
                risk="high",
                description=f"公式修复: {result.failed} 个公式转换失败"
            ))
    except ImportError:
        for para in doc.paragraphs:
            text = para.text
            if '$' in text and '\\frac' in text:
                items.append(CompatItem(
                    element_type="formula",
                    risk="medium",
                    description=f"检测到未转换公式: '{text[:30]}...'"
                ))

    return items
