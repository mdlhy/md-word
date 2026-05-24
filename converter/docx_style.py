"""DOCX style policy helpers for WPS-friendly Chinese documents."""

from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

DEFAULT_CN_FONT = "宋体"
DEFAULT_EN_FONT = "Times New Roman"
BLACK = "000000"

TEXT_STYLE_NAMES = {
    "Normal",
    "Body Text",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 4",
    "Heading 5",
    "Heading 6",
    "Title",
    "Subtitle",
    "List Paragraph",
    "List Bullet",
    "List Bullet 2",
    "List Bullet 3",
    "List Number",
    "List Number 2",
    "List Number 3",
    "Caption",
    "Quote",
    "Hyperlink",
}

CODE_STYLE_NAMES = {"Code", "Source Code"}


def ensure_rpr(element):
    rpr = element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        element.insert(0, rpr)
    return rpr


def set_rfonts(rpr, cn_font: str = DEFAULT_CN_FONT, en_font: str = DEFAULT_EN_FONT):
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)

    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    rfonts.set(qn("w:cs"), en_font)
    rfonts.set(qn("w:eastAsia"), cn_font)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        rfonts.attrib.pop(qn(f"w:{attr}"), None)


def force_black(rpr):
    color = rpr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rpr.append(color)
    color.set(qn("w:val"), BLACK)
    for attr in ("themeColor", "themeTint", "themeShade"):
        color.attrib.pop(qn(f"w:{attr}"), None)


def set_style_text_policy(style, cn_font: str = DEFAULT_CN_FONT, en_font: str = DEFAULT_EN_FONT):
    style.font.name = en_font
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    set_rfonts(rpr, cn_font, en_font)
    force_black(rpr)


def set_run_text_policy(run, cn_font: str = DEFAULT_CN_FONT, en_font: str = DEFAULT_EN_FONT):
    run.font.name = en_font
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = ensure_rpr(run._element)
    set_rfonts(rpr, cn_font, en_font)
    force_black(rpr)


def apply_document_text_policy(
    doc,
    cn_font: str = DEFAULT_CN_FONT,
    en_font: str = DEFAULT_EN_FONT,
):
    """Force regular document text to Songti/Times New Roman and black.

    WPS/Word built-in Heading styles often carry theme accent colors. Pandoc
    preserves those style definitions, so we normalize both styles and runs.
    Code styles are intentionally left alone to preserve monospaced snippets.
    """
    for style in doc.styles:
        if style.name in CODE_STYLE_NAMES:
            continue
        if style.name in TEXT_STYLE_NAMES or style.name.startswith("Heading"):
            set_style_text_policy(style, cn_font, en_font)

    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            if run.style and run.style.name in CODE_STYLE_NAMES:
                continue
            set_run_text_policy(run, cn_font, en_font)


def iter_all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)
